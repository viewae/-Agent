import logging

from docx import Document as DocxDocument

from app.parsers.base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    def extract_text(self, file_path: str) -> str:
        texts = []
        try:
            doc = DocxDocument(file_path)
            for para in doc.paragraphs:
                if para.text:
                    texts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    texts.append(row_text)
        except Exception:
            logger.exception("Failed to parse DOCX: %s", file_path)
            raise
        return "\n".join(texts)
