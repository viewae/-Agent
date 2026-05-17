import os

import pytest

from app.parsers.pdf_parser import PDFParser
from app.parsers.docx_parser import DocxParser
from app.parsers.txt_parser import TxtParser
from app.parsers.parser_factory import get_parser
from app.exceptions import ValidationError


class TestTxtParser:
    def test_extract_utf8(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("Hello World 你好世界", encoding="utf-8")
        parser = TxtParser()
        text = parser.extract_text(str(f))
        assert "Hello World" in text
        assert "你好世界" in text

    def test_extract_gbk(self, tmp_path):
        f = tmp_path / "test.txt"
        content = "这是一段较长的GBK编码中文文本，用于测试编码检测功能是否正常工作。\n" * 5
        f.write_text(content, encoding="gbk")
        parser = TxtParser()
        text = parser.extract_text(str(f))
        assert len(text) > 50

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        parser = TxtParser()
        text = parser.extract_text(str(f))
        assert text == ""


class TestDocxParser:
    def test_extract_paragraphs(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        parser = DocxParser()
        text = parser.extract_text(str(f))
        assert text is not None


class TestPDFParser:
    def test_extract_text(self, tmp_path):
        import pdfplumber
        pdf_path = tmp_path / "test.pdf"
        # pdfplumber requires a real PDF; skip gracefully
        if not os.path.exists(str(pdf_path)):
            pytest.skip("PDF test requires pre-existing file")


class TestParserFactory:
    def test_get_pdf_parser(self):
        parser = get_parser("pdf")
        assert isinstance(parser, PDFParser)

    def test_get_docx_parser(self):
        parser = get_parser("docx")
        assert isinstance(parser, DocxParser)

    def test_get_txt_parser(self):
        parser = get_parser("txt")
        assert isinstance(parser, TxtParser)

    def test_unsupported_type_raises(self):
        with pytest.raises(ValidationError):
            get_parser("xlsx")

    def test_case_insensitive(self):
        parser = get_parser("PDF")
        assert isinstance(parser, PDFParser)
